# ============================================================
# IMPORTS ORDENADOS
# ============================================================

import os
import re
import sys
import shutil
import subprocess
from datetime import datetime

from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
import pytesseract

import cv2
import numpy as np
from PIL import Image, ImageChops
from difflib import SequenceMatcher
import easyocr

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

import pandas as pd


# ============================================================
# DETECCIÓN DE MODO EXE Y RUTAS BASE
# ============================================================
def get_base_path():
    """Obtiene la ruta base (donde está el exe o el script)"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))

def get_resource_path(relative_path):
    """Obtiene ruta a recursos empaquetados en el exe"""
    if getattr(sys, 'frozen', False):
        # PyInstaller crea una carpeta temporal _MEIPASS
        base = sys._MEIPASS
    else:
        base = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base, relative_path)

BASE_PATH = get_base_path()


# ============================================================
# CONFIGURACIÓN (fácil de modificar)
# ============================================================
CONFIG = {
    "dpi": 150,
    "tolerancia_valor": 100,
    "similitud_minima": 0.85,
    "score_minimo_conexion": 2,
}

# Credenciales (mejor usar variables de entorno)
EMAIL_CONFIG = {
    "correo": os.environ.get("GMAIL_USER", "carterasuperla80@gmail.com"),
    "password": os.environ.get("GMAIL_APP_PASSWORD", "qdao uewp cnbt kraz"),
    "cc_cartera": "liquidacionla80@gmail.com",
}

# NITs que requieren edición manual antes de enviar (ej: Bucanero)
NITS_EDICION_MANUAL = [
    "800197463",  # Bucanero - requiere agregar descuentos
]


# ============================================================
# FUNCIONES AUXILIARES (fuera de ejecutar_unir)
# ============================================================

def safe_str(valor):
    return str(valor) if valor else ""


def limpiar_num(num):
    if not num:
        return ""
    try:
        return str(int(num))
    except ValueError:
        return ""


def similitud(a, b):
    if not a or not b:
        return 0
    return SequenceMatcher(None, a, b).ratio()


# ============================================================
# CARGAR RUTA_RUN
# ============================================================
def ejecutar_unir():

    RUTA_RUN = os.environ.get("RUTA_RUN")
    if not RUTA_RUN or not os.path.exists(RUTA_RUN):
        print("ERROR: No se recibió RUTA_RUN desde main.py o la ruta no existe.")
        sys.exit()

    # Rutas PDF
    pdf_comprobantes = os.path.join(RUTA_RUN, "Informe.pdf")
    pdf_pagos = os.path.join(RUTA_RUN, "Comprobante.pdf")

    # Carpetas internas
    CARPETAS = {
        "comprobantes": os.path.join(RUTA_RUN, "recortes_comprobantes"),
        "pagos": os.path.join(RUTA_RUN, "recortes_pagos"),
        "resultados": os.path.join(RUTA_RUN, "resultados"),
        "enviados": os.path.join(RUTA_RUN, "enviados"),
        "no_enviados": os.path.join(RUTA_RUN, "no_enviados"),
        "no_conectados": os.path.join(RUTA_RUN, "no_conectados"),
    }

    # Crear todas las carpetas
    for carpeta in CARPETAS.values():
        os.makedirs(carpeta, exist_ok=True)

    # ============================================================
    # CONFIGURAR TESSERACT (busca en exe empaquetado o instalado)
    # ============================================================
    tesseract_encontrado = False
    
    # Opción 1: Tesseract empaquetado con el exe
    tesseract_empaquetado = get_resource_path(os.path.join("tesseract", "tesseract.exe"))
    if os.path.exists(tesseract_empaquetado):
        pytesseract.pytesseract.tesseract_cmd = tesseract_empaquetado
        # Configurar TESSDATA_PREFIX para el exe
        tessdata_path = get_resource_path(os.path.join("tesseract", "tessdata"))
        os.environ["TESSDATA_PREFIX"] = tessdata_path
        print(f"Tesseract empaquetado: {tesseract_empaquetado}")
        tesseract_encontrado = True
    
    # Opción 2: Tesseract junto al exe (carpeta externa)
    if not tesseract_encontrado:
        tesseract_local = os.path.join(BASE_PATH, "tesseract", "tesseract.exe")
        if os.path.exists(tesseract_local):
            pytesseract.pytesseract.tesseract_cmd = tesseract_local
            os.environ["TESSDATA_PREFIX"] = os.path.join(BASE_PATH, "tesseract", "tessdata")
            print(f"Tesseract local: {tesseract_local}")
            tesseract_encontrado = True
    
    # Opción 3: Tesseract instalado en sistema
    if not tesseract_encontrado:
        try:
            subprocess.run(["tesseract", "--version"], check=True, 
                          stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            print("Tesseract detectado en PATH del sistema.")
            tesseract_encontrado = True
        except Exception:
            pass
    
    # Opción 4: Ruta típica de instalación Windows
    if not tesseract_encontrado:
        posible = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(posible):
            pytesseract.pytesseract.tesseract_cmd = posible
            print(f"Tesseract instalado: {posible}")
            tesseract_encontrado = True
    
    if not tesseract_encontrado:
        print("ERROR: No se encontró Tesseract. Coloca la carpeta 'tesseract' junto al exe.")
        sys.exit()

    # ============================================================
    # CONFIGURAR POPPLER (para pdf2image)
    # ============================================================
    poppler_path = None
    
    # Opción 1: Poppler empaquetado
    poppler_empaquetado = get_resource_path(os.path.join("poppler", "bin"))
    if os.path.exists(poppler_empaquetado):
        poppler_path = poppler_empaquetado
        print(f"Poppler empaquetado: {poppler_path}")
    
    # Opción 2: Poppler junto al exe
    if not poppler_path:
        poppler_local = os.path.join(BASE_PATH, "poppler", "bin")
        if os.path.exists(poppler_local):
            poppler_path = poppler_local
            print(f"Poppler local: {poppler_path}")
    
    # Opción 3: Poppler en PATH (no hace falta especificar)
    if not poppler_path:
        try:
            subprocess.run(["pdftoppm", "-v"], check=True, 
                          stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            print("Poppler detectado en PATH del sistema.")
        except Exception:
            print("ADVERTENCIA: Poppler no encontrado. Coloca la carpeta 'poppler' junto al exe.")

    # ============================================================
    # CONFIGURAR EASYOCR (lazy loading con modelos locales)
    # ============================================================
    _easy_reader = None
    
    # Configurar carpeta de modelos EasyOCR junto al exe
    easyocr_model_dir = os.path.join(BASE_PATH, "easyocr_models")
    if not os.path.exists(easyocr_model_dir):
        # También buscar en recursos empaquetados
        easyocr_model_dir_packed = get_resource_path("easyocr_models")
        if os.path.exists(easyocr_model_dir_packed):
            easyocr_model_dir = easyocr_model_dir_packed
        else:
            os.makedirs(easyocr_model_dir, exist_ok=True)

    def get_easy_reader():
        nonlocal _easy_reader
        if _easy_reader is None:
            print("Inicializando EasyOCR...")
            print(f"Modelos EasyOCR en: {easyocr_model_dir}")
            _easy_reader = easyocr.Reader(
                ['es'], 
                gpu=False,
                model_storage_directory=easyocr_model_dir,
                download_enabled=True  # Descarga si no existen
            )
        return _easy_reader

    # ============================================================
    # VALIDACIÓN DE PDFs
    # ============================================================
    if not os.path.exists(pdf_comprobantes) or not os.path.exists(pdf_pagos):
        print("ERROR: No se encontraron los archivos requeridos:")
        print(" - " + pdf_comprobantes)
        print(" - " + pdf_pagos)
        sys.exit()

    # ============================================================
    # FUNCIONES DE IMAGEN
    # ============================================================
    def recortar_bordes(img):
        fondo = Image.new(img.mode, img.size, img.getpixel((0, 0)))
        dif = ImageChops.difference(img, fondo)
        bbox = dif.getbbox()
        return img.crop(bbox) if bbox else img

    def dividir_pagina_en_pagos(imagen, salida_base):
        gray = cv2.cvtColor(np.array(imagen), cv2.COLOR_RGB2GRAY)
        _, thresh = cv2.threshold(gray, 240, 255, cv2.THRESH_BINARY_INV)

        projection = np.sum(thresh, axis=1)
        limites, dentro, inicio = [], False, 0

        for i, val in enumerate(projection):
            if val > 700 and not dentro:
                inicio, dentro = i, True
            elif val < 700 and dentro:
                limites.append((inicio, i))
                dentro = False

        recortes = []
        for j, (y1, y2) in enumerate(limites):
            if y2 - y1 > 90:
                crop = imagen.crop((0, y1, imagen.width, y2))
                path = f"{salida_base}_{j+1}.jpg"
                crop.save(path)
                recortes.append(path)

        return recortes

    def procesar_pdf(pdf, carpeta, tipo, dividir=False):
        print(f"Convirtiendo PDF: {tipo}...")
        # Usar poppler_path si está configurado
        if poppler_path:
            paginas = convert_from_path(pdf, dpi=CONFIG["dpi"], poppler_path=poppler_path)
        else:
            paginas = convert_from_path(pdf, dpi=CONFIG["dpi"])
        rutas = []

        for i, img in enumerate(paginas):
            img = recortar_bordes(img)
            if dividir:
                recortes = dividir_pagina_en_pagos(img, os.path.join(carpeta, f"{tipo}_{i+1}"))
                rutas.extend(recortes)
            else:
                ruta = os.path.join(carpeta, f"{tipo}_{i+1}.jpg")
                img.save(ruta)
                rutas.append(ruta)

        print(f"{len(rutas)} imagenes generadas para {tipo}")
        return rutas

    # ============================================================
    # OCR UNIFICADO (una sola función con modos)
    # ============================================================
    def leer_texto_ocr(imagen, modo="normal"):
        """
        Modos:
        - "normal": OCR básico con Tesseract
        - "mejorado": OCR con preprocesamiento avanzado
        - "easyocr": Usa EasyOCR
        """
        if modo == "easyocr":
            reader = get_easy_reader()
            resultados = reader.readtext(imagen, detail=0, paragraph=True)
            return " ".join(resultados).upper().replace("\n", " ")

        img = cv2.imread(imagen, cv2.IMREAD_GRAYSCALE)
        
        if modo == "mejorado":
            img = cv2.resize(img, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
            img = cv2.GaussianBlur(img, (3, 3), 0)
            img = cv2.adaptiveThreshold(img, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                        cv2.THRESH_BINARY, 31, 2)
            config = "--psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        else:
            img = cv2.threshold(img, 180, 255, cv2.THRESH_BINARY)[1]
            config = ""

        texto = pytesseract.image_to_string(img, lang="spa", config=config)
        return texto.upper().replace("\n", " ")

    # ============================================================
    # EXTRACCIÓN DE DATOS
    # ============================================================
    def extraer_datos(texto, es_pago=False):
        nombre, numero, valor = None, None, None

        # --- EXTRAER NOMBRE ---
        m_nombre = re.search(r"(?:BENEFICIARIO|NOMBRE DE BENEFICIARIO)[:\.\s]+([A-ZÑ&.\s]{10,})", texto)
        if not m_nombre:
            m_nombre = re.search(
                r"NÚMERO DE BENEFICIARIO[:\.\s]+([A-ZÑ&.\s]+?)(?:\s+NIT|\s+DOCUMENTO|\s+CÉDULA|$)",
                texto
            )
        if not m_nombre:
            m_nombre = re.search(
                r"([A-ZÑ\s]{10,})\s+(NIT|DOCUMENTO|CÉDULA|NÚMERO DE BENEFICIARIO)",
                texto
            )

        if m_nombre:
            nombre = m_nombre.group(1).strip()
            nombre = re.sub(r"\s{2,}", " ", nombre)
            nombre = re.sub(r"(?i)\s+(NIT|DOCUMENTO|CÉDULA).*", "", nombre)

        # --- EXTRAER NÚMERO (NIT/DOCUMENTO) ---
        if es_pago:
            m_doc = re.search(r"DOCUMENTO\s*[:\.\s]*([0-9\.\-]+)", texto)
            if m_doc:
                numero = re.sub(r"\D", "", m_doc.group(1))
        else:
            # Para el informe/comprobante de egreso, buscar específicamente
            # el NIT que está en la sección del proveedor
            # El formato típico es: "NIT: 71690934 -1" (donde -1 es el dígito de verificación a ignorar)
            
            m_nit = None
            
            # Patrón 1: NIT seguido de número, capturando el NIT principal y el dígito verificador por separado
            # Busca el patrón "NIT: XXXXXXXX -X" o "NIT: XXXXXXXXX" cerca de CIUDAD
            m_nit = re.search(r"NIT[:\s]*(\d{7,10})[\s\-]*(\d)?\s*(?:CIUDAD|VALOR|$)", texto, re.IGNORECASE)
            
            if not m_nit:
                # Patrón 2: NIT con puntos (ej: 71.690.934-1)
                m_nit = re.search(r"NIT[:\s]*([\d\.]{8,15})[\s\-]*(\d)?", texto, re.IGNORECASE)
            
            if not m_nit:
                # Patrón 3: Más flexible - NIT seguido de cualquier número
                m_nit = re.search(r"NIT[:\s\.]*(\d[\d\.]+)", texto, re.IGNORECASE)
            
            if m_nit:
                numero_raw = m_nit.group(1)
                # Eliminar puntos y obtener solo dígitos
                numero = re.sub(r"\D", "", numero_raw)
                
                # Si el número tiene más de 11 dígitos, probablemente tiene basura
                # Los NITs colombianos tienen entre 8 y 10 dígitos (+ 1 de verificador máximo = 11)
                if len(numero) > 11:
                    numero = numero[:10]  # Tomar solo los primeros 10 dígitos
            
            # Si aún no se encontró, usar patrones de respaldo
            if not numero:
                patrones_nit = [
                    r"(?:DOCUMENTO|CÉDULA|NÚMERO DE BENEFICIARIO)[:\.\s]*([0-9\.\-]+)",
                ]
                for patron in patrones_nit:
                    m_nit = re.search(patron, texto)
                    if m_nit:
                        numero_raw = m_nit.group(1) if m_nit.groups() else m_nit.group(0)
                        if "-" in numero_raw:
                            numero_raw = numero_raw.split("-")[0]
                        numero = re.sub(r"\D", "", numero_raw)
                        if len(numero) >= 8:
                            break

        if numero:
            numero = numero.lstrip("0")

        # --- EXTRAER VALOR ---
        m_valor = re.search(r"\$?\s*\d{1,3}(?:\.\d{3})+(?:,\d{2})?", texto)
        if m_valor:
            valor_raw = m_valor.group(0).replace("$", "").replace(" ", "")
            valor = valor_raw.replace(".", "").replace(",", "")

        return nombre, numero, valor

    # ============================================================
    # PROCESAR CARPETAS (OCR inicial)
    # ============================================================
    def procesar_carpeta(rutas, tipo):
        datos = []
        print(f"Analizando {tipo}s...")
        
        for r in rutas:
            texto = leer_texto_ocr(r, modo="normal")
            
            # DEBUG: Mostrar parte del texto OCR para comprobantes
            if tipo == "comprobante":
                if "NIT" in texto.upper():
                    idx = texto.upper().find("NIT")
                    print(f"[DEBUG OCR] Texto alrededor de NIT: ...{texto[max(0,idx-20):idx+50]}...")
                else:
                    print(f"[DEBUG OCR] No se encontró 'NIT' en el texto. Primeros 200 chars: {texto[:200]}")
            
            nombre, numero, valor = extraer_datos(texto, es_pago=(tipo == "pago"))
            
            # Si es comprobante y NO se encontró el NIT, intentar con otros métodos OCR
            if tipo == "comprobante" and not numero:
                print(f"[REINTENTO] {os.path.basename(r)} - Intentando OCR mejorado...")
                texto_mejorado = leer_texto_ocr(r, modo="mejorado")
                nombre2, numero2, valor2 = extraer_datos(texto_mejorado, es_pago=False)
                
                if numero2:
                    numero = numero2
                    nombre = nombre or nombre2
                    valor = valor or valor2
                    print(f"[REINTENTO OK] Encontrado con OCR mejorado: NIT {numero}")
                else:
                    # Último intento: EasyOCR
                    print(f"[REINTENTO] {os.path.basename(r)} - Intentando EasyOCR...")
                    texto_easy = leer_texto_ocr(r, modo="easyocr")
                    nombre3, numero3, valor3 = extraer_datos(texto_easy, es_pago=False)
                    
                    if numero3:
                        numero = numero3
                        nombre = nombre or nombre3
                        valor = valor or valor3
                        print(f"[REINTENTO OK] Encontrado con EasyOCR: NIT {numero}")
                    else:
                        print(f"[REINTENTO FALLIDO] No se pudo extraer NIT de {os.path.basename(r)}")
            
            datos.append({"archivo": r, "nombre": nombre, "numero": numero, "valor": valor})
            print(f"{os.path.basename(r)} -> {nombre or 'Sin nombre'} | NIT: {safe_str(numero)} | Valor: {safe_str(valor)}")

        return datos

    # Procesar PDFs
    rutas_comprobantes = procesar_pdf(pdf_comprobantes, CARPETAS["comprobantes"], "comprobante", dividir=False)
    rutas_pagos = procesar_pdf(pdf_pagos, CARPETAS["pagos"], "pago", dividir=True)

    comprobantes = procesar_carpeta(rutas_comprobantes, "comprobante")
    pagos = procesar_carpeta(rutas_pagos, "pago")

    # ============================================================
    # AGRUPAR COMPROBANTES POR NIT
    # ============================================================
    def agrupar_por_nit(lista):
        grupos = {}
        for item in lista:
            nit = item["numero"] or "SIN_NIT"
            grupos.setdefault(nit, []).append(item)
        return grupos

    grupos_comprobantes = agrupar_por_nit(comprobantes)

    # ============================================================
    # CALCULAR SCORE DE COINCIDENCIA
    # ============================================================
    def nits_coinciden(nit1, nit2):
        """Verifica si dos NITs coinciden, considerando posibles errores de OCR"""
        if not nit1 or not nit2:
            return False
        
        n1 = limpiar_num(nit1)
        n2 = limpiar_num(nit2)
        
        # Coincidencia exacta
        if n1 == n2:
            return True
        
        # Si la diferencia de longitud es muy grande, no son el mismo NIT
        if abs(len(n1) - len(n2)) > 2:
            return False
        
        # Verificar si uno contiene al otro
        # Esto cubre los casos:
        # - 171690934 vs 71690934 (dígito extra al inicio por ruido OCR)
        # - 8902018814 vs 890201881 (dígito verificador pegado al final)
        # - 716909341 vs 71690934 (dígito extra al final)
        if n1 in n2 or n2 in n1:
            return True
        
        return False
    
    def nit_del_pago_coincide_con_informe(nit_pago, nit_informe):
        """
        Verifica si el NIT del pago (que es el correcto) coincide con el del informe.
        El NIT del pago es la referencia porque siempre se extrae bien.
        El NIT del informe puede tener errores de OCR.
        """
        if not nit_pago or not nit_informe:
            return False
        
        pago = limpiar_num(nit_pago)
        informe = limpiar_num(nit_informe)
        
        # Coincidencia exacta
        if pago == informe:
            return True
        
        # El NIT del pago debe estar contenido en el del informe
        # (por si el informe tiene dígitos extra por error OCR)
        if pago in informe:
            return True
        
        # O el del informe empieza con el del pago (dígito verificador pegado)
        if informe.startswith(pago):
            return True
        
        # O el del informe termina con el del pago (ruido al inicio)
        if informe.endswith(pago):
            return True
        
        return False
    
    def calcular_score(pago, comp):
        """Calcula score de coincidencia entre un pago y un comprobante"""
        score = 0

        # 1️⃣ Documento / NIT - El NIT del pago es la referencia correcta
        if pago["numero"] and comp["numero"]:
            if nit_del_pago_coincide_con_informe(pago["numero"], comp["numero"]):
                score += 1

        # 2️⃣ Valor con tolerancia
        if pago["valor"] and comp["valor"]:
            try:
                if abs(int(pago["valor"]) - int(comp["valor"])) <= CONFIG["tolerancia_valor"]:
                    score += 1
            except ValueError:
                pass

        # 3️⃣ Nombre similar
        if similitud(pago["nombre"], comp["nombre"]) >= CONFIG["similitud_minima"]:
            score += 1

        return score

    # ============================================================
    # SEGUNDO INTENTO DE CONEXIÓN (CORREGIDO)
    # ============================================================
    def segundo_intento_conexion(pagos_sobrantes, comprobantes_no_conectados):
        """
        CORRECCIÓN: No modifica los datos originales del pago.
        Usa variables locales para el OCR mejorado.
        """
        nuevas_conexiones = []
        comprobantes_usados = set()

        for pago in pagos_sobrantes:
            # Variables LOCALES para este pago (no modificamos el original aún)
            nom_p, num_p, val_p = pago["nombre"], pago["numero"], pago["valor"]

            # 🔁 Intento 1: OCR mejorado con Tesseract
            texto_mejorado = leer_texto_ocr(pago["archivo"], modo="mejorado")
            nom_m, num_m, val_m = extraer_datos(texto_mejorado, es_pago=True)

            # Usar datos mejorados si los anteriores estaban vacíos
            nom_p = nom_p or nom_m
            num_p = num_p or num_m
            val_p = val_p or val_m

            # 🔁 Intento 2: EasyOCR SOLO si aún no hay número
            if not num_p:
                texto_easy = leer_texto_ocr(pago["archivo"], modo="easyocr")
                nom_e, num_e, val_e = extraer_datos(texto_easy, es_pago=True)
                nom_p = nom_p or nom_e
                num_p = num_p or num_e
                val_p = val_p or val_e

            # Crear diccionario temporal con datos mejorados para comparar
            pago_mejorado = {"nombre": nom_p, "numero": num_p, "valor": val_p}

            # Buscar mejor coincidencia
            mejor_comp = None
            mejor_score = 0

            for comp in comprobantes_no_conectados:
                if id(comp) in comprobantes_usados:
                    continue

                score = calcular_score(pago_mejorado, comp)

                if score >= CONFIG["score_minimo_conexion"] and score > mejor_score:
                    mejor_score = score
                    mejor_comp = comp

            # Si encontramos coincidencia, actualizar pago y registrar
            if mejor_comp:
                pago["nombre"] = nom_p
                pago["numero"] = num_p
                pago["valor"] = val_p
                comprobantes_usados.add(id(mejor_comp))
                nuevas_conexiones.append((mejor_comp, pago))

        return nuevas_conexiones

    # ============================================================
    # EMPAREJAR GRUPOS DE COMPROBANTES CON PAGOS
    # ============================================================
    print("Emparejando comprobantes agrupados por NIT...")

    pares = []
    pagos_disponibles = pagos.copy()  # Copia para no modificar original

    for nit, lista_comps in grupos_comprobantes.items():
        mejor_pago = None

        for pago in pagos_disponibles:
            # El NIT del pago es el correcto, verificar si coincide con el del informe
            if nit_del_pago_coincide_con_informe(pago["numero"], nit):
                mejor_pago = pago
                break

        if mejor_pago:
            pagos_disponibles.remove(mejor_pago)

        pares.append((nit, lista_comps, mejor_pago))

    # ============================================================
    # FUNCIÓN PARA CREAR WORD
    # ============================================================
    def crear_word(nit, lista_comps, pago, carpeta_destino):
        """Crea un documento Word para un NIT dado"""
        nombre_archivo = re.sub(r'[\\\/*?:"<>|]', "", str(nit))

        doc = Document()
        doc.add_heading(f"Comprobantes agrupados — NIT {nit}", level=1)

        for comp in lista_comps:
            doc.add_paragraph(f"Página del comprobante — {comp['nombre']}")
            doc.add_paragraph(f"NIT: {safe_str(comp['numero'])} | Valor: {safe_str(comp['valor'])}")
            doc.add_picture(comp["archivo"], width=Inches(6.5))
            doc.add_page_break()

        if pago:
            doc.add_heading("PAGO CORRESPONDIENTE", level=2)
            doc.add_paragraph(f"NIT: {safe_str(pago['numero'])} | Valor: {safe_str(pago['valor'])}")
            doc.add_picture(pago["archivo"], width=Inches(6.5))
        else:
            doc.add_paragraph("NO SE ENCONTRÓ PAGO.")

        path = os.path.join(carpeta_destino, f"{nombre_archivo}.docx")
        doc.save(path)
        print(f"Guardado: {path}")
        return path

    # Crear Word SOLO para pares que tienen pago
    for nit, lista_comps, pago in pares:
        if pago:  # Solo si hay pago asociado
            crear_word(nit, lista_comps, pago, CARPETAS["resultados"])


    # ============================================================
    # RESUMEN NO CONECTADOS
    # ============================================================
    comprobantes_no_conectados = []
    for nit, lista_comps, pago in pares:
        if pago is None:
            comprobantes_no_conectados.extend(lista_comps)

    pagos_sobrantes = pagos_disponibles.copy()

    # ============================================================
    # SEGUNDO ESCANEO AUTOMÁTICO
    # ============================================================
    print("🔁 Iniciando segundo intento de conexión inteligente...")

    nuevas_conexiones = segundo_intento_conexion(pagos_sobrantes, comprobantes_no_conectados)

    # Crear Word para las nuevas conexiones del segundo intento
    for comp, pago in nuevas_conexiones:
        print(f"✔ Conectado en segundo intento: {comp['numero']} ↔ {pago['numero']}")

        # Crear el Word para esta nueva conexión
        crear_word(comp["numero"], [comp], pago, CARPETAS["resultados"])

        # Remover de las listas de no conectados
        if pago in pagos_sobrantes:
            pagos_sobrantes.remove(pago)
        if comp in comprobantes_no_conectados:
            comprobantes_no_conectados.remove(comp)

    # ============================================================
    # REPORTE DE NO CONECTADOS
    # ============================================================
    if comprobantes_no_conectados or pagos_sobrantes:
        resumen = Document()
        resumen.add_heading("REVISION MANUAL REQUERIDA", level=1)
        resumen.add_paragraph("Comprobantes sin pago o pagos sobrantes.\n")

        if comprobantes_no_conectados:
            resumen.add_heading("COMPROBANTES SIN PAGO", level=2)
            for c in comprobantes_no_conectados:
                resumen.add_paragraph(f"- {c['nombre'] or 'Sin nombre'} | NIT: {safe_str(c['numero'])} | Valor: {safe_str(c['valor'])}")
                run = resumen.add_paragraph().add_run()
                run.add_picture(c["archivo"], width=Inches(6.5))
            resumen.add_page_break()

        if pagos_sobrantes:
            resumen.add_heading("PAGOS SOBRANTES", level=2)
            for pago in pagos_sobrantes:
                resumen.add_paragraph(f"NIT: {safe_str(pago['numero'])} | Valor: {safe_str(pago['valor'])}")
                run = resumen.add_paragraph().add_run()
                run.add_picture(pago["archivo"], width=Inches(6.5))

        resumen.save(os.path.join(CARPETAS["no_conectados"], "NO_CONECTADOS_Y_PAGOS_SOBRANTES.docx"))
        print(f"Guardado reporte de no conectados: {CARPETAS['no_conectados']}")

    # ============================================================    
    # ENVÍO AUTOMÁTICO GMAIL (conexión reutilizada)
    # ============================================================
    fallidos = []

    try:
        print("Iniciando envio por Gmail...")

        excel_path = os.path.join(RUTA_RUN, "proveedores_correos.xlsx")
        if not os.path.exists(excel_path):
            excel_path = "proveedores_correos.xlsx"

        if not os.path.exists(excel_path):
            print("ERROR: No encontre proveedores_correos.xlsx")
        else:
            df = pd.read_excel(excel_path, dtype=str)
            df["NIT"] = df["NIT"].astype(str).str.replace(r"\D", "", regex=True).str.strip()

            nit_to_email = dict(zip(df["NIT"], df.get("CORREO", "")))
            nit_to_name = dict(zip(df["NIT"], df.get("NOMBRE_PROVEEDOR", "")))

            enviados = 0
            archivos_a_enviar = [f for f in os.listdir(CARPETAS["resultados"]) if f.endswith(".docx")]

            # Una sola conexión SMTP para todos los correos
            if archivos_a_enviar:
                try:
                    server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
                    server.login(EMAIL_CONFIG["correo"], EMAIL_CONFIG["password"])

                    for archivo in archivos_a_enviar:
                        nit = os.path.splitext(archivo)[0]
                        correos_raw = nit_to_email.get(nit, "") or ""
                        correos_raw = str(correos_raw) if correos_raw and str(correos_raw).lower() != "nan" else ""
                        lista_correos = [c.strip() for c in correos_raw.split(";") if c.strip()]

                        if not lista_correos:
                            print(f"FALLIDO: {nit} -> Sin correos en Excel")
                            fallidos.append({"nit": nit, "archivo": archivo, "motivo": "Sin correo"})
                            continue

                        nombre = nit_to_name.get(nit, "Proveedor")
                        
                        ruta_archivo = os.path.join(CARPETAS["resultados"], archivo)
                        
                        # Si es un NIT que requiere edición manual, abrir Word y esperar
                        if nit in NITS_EDICION_MANUAL:
                            print(f"\n⚠️ NIT {nit} ({nombre}) requiere edición manual.")
                            print(f"Abriendo documento para editar: {ruta_archivo}")
                            os.startfile(ruta_archivo)
                            
                            # Importar messagebox aquí para no afectar imports globales
                            from tkinter import messagebox
                            respuesta = messagebox.askyesno(
                                "Edición Manual Requerida",
                                f"Se abrió el documento de {nombre} (NIT {nit}).\n\n"
                                f"1. Edita el Word agregando descuentos u otros datos\n"
                                f"2. GUARDA el documento (Ctrl+S)\n"
                                f"3. Cierra el Word\n\n"
                                f"¿Ya terminaste de editar y guardaste el documento?"
                            )
                            
                            if not respuesta:
                                print(f"[SALTADO] {nit} - Usuario canceló la edición")
                                fallidos.append({"nit": nit, "archivo": archivo, "motivo": "Edición cancelada por usuario"})
                                continue
                        
                        body = f"""Buenos días,

Adjuntamos el comprobante de pago correspondiente al NIT {nit}.

Quedamos atentos.
Cordialmente,
Equipo de Cartera"""

                        try:
                            msg = MIMEMultipart()
                            msg['From'] = EMAIL_CONFIG["correo"]
                            msg['To'] = ", ".join(lista_correos)
                            msg['Cc'] = EMAIL_CONFIG["cc_cartera"]
                            msg['Subject'] = f"Comprobante de pago - {nombre} - NIT {nit}"
                            msg.attach(MIMEText(body, 'plain'))

                            with open(ruta_archivo, "rb") as adjunto:
                                part = MIMEBase('application', 'octet-stream')
                                part.set_payload(adjunto.read())
                                encoders.encode_base64(part)
                                part.add_header('Content-Disposition', f'attachment; filename="{archivo}"')
                                msg.attach(part)

                            destinatarios = lista_correos + [EMAIL_CONFIG["cc_cartera"]]
                            server.sendmail(EMAIL_CONFIG["correo"], destinatarios, msg.as_string())

                            print(f"Enviado correctamente a {lista_correos} con copia a cartera")
                            enviados += 1
                            shutil.copy(ruta_archivo, os.path.join(CARPETAS["enviados"], archivo))

                        except Exception as e:
                            fallidos.append({"nit": nit, "archivo": archivo, "motivo": str(e)})
                            print(f"ERROR enviando a {nit}: {e}")

                    server.quit()

                except smtplib.SMTPAuthenticationError:
                    print("ERROR: Credenciales de Gmail incorrectas")
                except Exception as e:
                    print(f"ERROR conectando a Gmail: {e}")

            print(f"Envio terminado. Enviados correctamente: {enviados}")

    except Exception as e:
        print(f"ERROR general en el envio: {e}")

    # ============================================================
    # MOVER LOS QUE NO SE ENVIARON
    # ============================================================
    if fallidos:
        for f in fallidos:
            origen = os.path.join(CARPETAS["resultados"], f["archivo"])
            destino = os.path.join(CARPETAS["no_enviados"], f["archivo"])
            if os.path.exists(origen):
                shutil.copy(origen, destino)
                print(f"No enviado -> movido a no_enviados: {f['archivo']}")

        print("Algunos comprobantes no se enviaron. Revisa la carpeta 'no_enviados'.")
    else:
        print("Todos los comprobantes fueron enviados correctamente.")

